from docx import Document
from docx.shared import Pt, Inches, RGBColor, Emu
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from lxml import etree
import os
import io


def make_image_floating(run, width_emu, height_emu, pos_x_emu, pos_v_emu):
    """
    Convert inline image to floating anchor.
    Page-anchored, In Front of Text wrapping.
    """
    inline = run._r.find('.//{http://schemas.openxmlformats.org/drawingml/2006/wordprocessingDrawing}inline')
    if inline is None:
        return
    
    drawing = inline.getparent()
    
    nsmap = {
        'wp': 'http://schemas.openxmlformats.org/drawingml/2006/wordprocessingDrawing',
        'a': 'http://schemas.openxmlformats.org/drawingml/2006/main',
        'r': 'http://schemas.openxmlformats.org/officeDocument/2006/relationships',
        'pic': 'http://schemas.openxmlformats.org/drawingml/2006/picture',
    }
    
    anchor = etree.Element('{http://schemas.openxmlformats.org/drawingml/2006/wordprocessingDrawing}anchor', nsmap=nsmap)
    anchor.set('distT', '0')
    anchor.set('distB', '0')
    anchor.set('distL', '0')
    anchor.set('distR', '0')
    anchor.set('simplePos', '0')
    anchor.set('relativeHeight', '251658240')
    anchor.set('behindDoc', '0')
    anchor.set('locked', '0')
    anchor.set('layoutInCell', '1')
    anchor.set('allowOverlap', '1')
    
    simple_pos = etree.SubElement(anchor, '{http://schemas.openxmlformats.org/drawingml/2006/wordprocessingDrawing}simplePos')
    simple_pos.set('x', '0')
    simple_pos.set('y', '0')
    
    pos_h = etree.SubElement(anchor, '{http://schemas.openxmlformats.org/drawingml/2006/wordprocessingDrawing}positionH')
    pos_h.set('relativeFrom', 'page')
    pos_offset_h = etree.SubElement(pos_h, '{http://schemas.openxmlformats.org/drawingml/2006/wordprocessingDrawing}posOffset')
    pos_offset_h.text = str(int(pos_x_emu))
    
    pos_v = etree.SubElement(anchor, '{http://schemas.openxmlformats.org/drawingml/2006/wordprocessingDrawing}positionV')
    pos_v.set('relativeFrom', 'page')
    pos_offset_v = etree.SubElement(pos_v, '{http://schemas.openxmlformats.org/drawingml/2006/wordprocessingDrawing}posOffset')
    pos_offset_v.text = str(int(pos_v_emu))
    
    extent_elem = etree.SubElement(anchor, '{http://schemas.openxmlformats.org/drawingml/2006/wordprocessingDrawing}extent')
    extent_elem.set('cx', str(int(width_emu)))
    extent_elem.set('cy', str(int(height_emu)))
    
    effect = etree.SubElement(anchor, '{http://schemas.openxmlformats.org/drawingml/2006/wordprocessingDrawing}effectExtent')
    effect.set('l', '0')
    effect.set('t', '0')
    effect.set('r', '0')
    effect.set('b', '0')
    
    etree.SubElement(anchor, '{http://schemas.openxmlformats.org/drawingml/2006/wordprocessingDrawing}wrapNone')
    
    for child in list(inline):
        anchor.append(child)
    
    drawing.remove(inline)
    drawing.append(anchor)


def add_border_to_cell(cell):
    """Add borders to table cell"""
    tc = cell._element
    tcPr = tc.get_or_add_tcPr()
    tcBorders = OxmlElement('w:tcBorders')
    for edge in ('top', 'left', 'bottom', 'right'):
        edge_element = OxmlElement(f'w:{edge}')
        edge_element.set(qn('w:val'), 'single')
        edge_element.set(qn('w:sz'), '4')
        edge_element.set(qn('w:space'), '0')
        edge_element.set(qn('w:color'), '000000')
        tcBorders.append(edge_element)
    tcPr.append(tcBorders)


def set_cell_width(cell, width_inches):
    """Set cell width"""
    tc = cell._element
    tcPr = tc.get_or_add_tcPr()
    tcW = OxmlElement('w:tcW')
    tcW.set(qn('w:w'), str(int(width_inches * 1440)))
    tcW.set(qn('w:type'), 'dxa')
    tcPr.append(tcW)


def set_column_widths(table, widths_inches):
    """Set column widths for a table"""
    for row in table.rows:
        for idx, cell in enumerate(row.cells):
            if idx < len(widths_inches):
                set_cell_width(cell, widths_inches[idx])


def remove_table_borders(table):
    """Remove all borders from a table"""
    tbl = table._tbl
    tblPr = tbl.tblPr if tbl.tblPr is not None else OxmlElement('w:tblPr')
    tblBorders = OxmlElement('w:tblBorders')
    for edge in ('top', 'left', 'bottom', 'right', 'insideH', 'insideV'):
        edge_element = OxmlElement(f'w:{edge}')
        edge_element.set(qn('w:val'), 'nil')
        tblBorders.append(edge_element)
    tblPr.append(tblBorders)
    if tbl.tblPr is None:
        tbl.insert(0, tblPr)


def create_exam_paper(
    department="CSBS",
    section_val="A",
    semester="IV",
    date_session="2025-02-27 (FN)",
    course_code="CBB1222",
    course_name="Operating Systems",
    cia_type="CIA-I",
    qp_type="QP-I",
    part_a_questions=None,
    part_b_questions=None,
    part_c_questions=None,
    output_path=None
):
    """
    Create CIA exam paper matching CIA_ACTUAL_FORMAT.docx exactly.
    Header: Left logo + text, Right logo (floating)
    REG No: LEFT-aligned
    """
    
    is_qp_type_2 = qp_type == "QP-II"
    part_b_marks_each = 16 if is_qp_type_2 else 12
    part_b_count = 2 if is_qp_type_2 else 4
    
    # Default questions
    if part_a_questions is None:
        co1 = 'CO1' if cia_type == 'CIA-I' else 'CO3'
        co2 = 'CO2' if cia_type == 'CIA-I' else 'CO4'
        part_a_questions = [
            {'qno': '1.', 'question': 'Define real time systems.', 'co': co1, 'btl': 'BTL1', 'marks': '2'},
            {'qno': '2.', 'question': 'Describe Paravirtualization.', 'co': co1, 'btl': 'BTL1', 'marks': '2'},
            {'qno': '3.', 'question': 'Identify the main advantages of virtual Machine.', 'co': co1, 'btl': 'BTL3', 'marks': '2'},
            {'qno': '4.', 'question': 'How can starvation of processes be avoided in priority scheduling?', 'co': co2, 'btl': 'BTL3', 'marks': '2'},
            {'qno': '5.', 'question': 'Draw and explain data fields associated with process control blocks.', 'co': co2, 'btl': 'BTL2', 'marks': '2'},
            {'qno': '6.', 'question': 'Define the term dispatch latency.', 'co': co2, 'btl': 'BTL1', 'marks': '2'}
        ]
    
    if part_b_questions is None:
        co1 = 'CO1' if cia_type == 'CIA-I' else 'CO3'
        co2 = 'CO2' if cia_type == 'CIA-I' else 'CO4'
        marks = str(part_b_marks_each)
        
        if is_qp_type_2:
            part_b_questions = [
                {'qno': '7.(a)', 'question': 'Classify Interrupts in an Operating System and Explain Each Type with Examples.', 'co': co1, 'btl': 'BTL3', 'marks': marks},
                {'qno': '(OR)', 'question': '(OR)', 'co': '', 'btl': '', 'marks': ''},
                {'qno': '7.(b)', 'question': 'Explain the different operating system structures with neat sketch.', 'co': co1, 'btl': 'BTL4', 'marks': marks},
                {'qno': '8.(a)', 'question': 'Explain the Resource Manager View of an operating system.', 'co': co2, 'btl': 'BTL4', 'marks': marks},
                {'qno': '(OR)', 'question': '(OR)', 'co': '', 'btl': '', 'marks': ''},
                {'qno': '8.(b)', 'question': 'Construct a short note on basic architectural concepts of OS.', 'co': co2, 'btl': 'BTL3', 'marks': marks}
            ]
        else:
            part_b_questions = [
                {'qno': '7.(a)', 'question': 'Classify Interrupts in an Operating System and Explain Each Type with Examples.', 'co': co1, 'btl': 'BTL2', 'marks': marks},
                {'qno': '(OR)', 'question': '(OR)', 'co': '', 'btl': '', 'marks': ''},
                {'qno': '7.(b)', 'question': 'Explain the different operating system structures with neat sketch.', 'co': co1, 'btl': 'BTL2', 'marks': marks},
                {'qno': '8.(a)', 'question': 'Explain the Resource Manager View of an operating system.', 'co': co1, 'btl': 'BTL3', 'marks': marks},
                {'qno': '(OR)', 'question': '(OR)', 'co': '', 'btl': '', 'marks': ''},
                {'qno': '8.(b)', 'question': 'Construct a short note on basic architectural concepts of OS.', 'co': co1, 'btl': 'BTL3', 'marks': marks},
                {'qno': '9.(a)', 'question': 'Summarize components of process and various states of a process.', 'co': co2, 'btl': 'BTL2', 'marks': marks},
                {'qno': '(OR)', 'question': '(OR)', 'co': '', 'btl': '', 'marks': ''},
                {'qno': '9.(b)', 'question': 'Explain about the different types of scheduler and its scheduling criteria.', 'co': co2, 'btl': 'BTL2', 'marks': marks},
                {'qno': '10.(a)', 'question': 'Consider the following set of processes with CPU burst time. Draw Gantt chart and find average turnaround time.', 'co': co2, 'btl': 'BTL3', 'marks': marks},
                {'qno': '(OR)', 'question': '(OR)', 'co': '', 'btl': '', 'marks': ''},
                {'qno': '10.(b)', 'question': 'Calculate average waiting time and turnaround time using FCFS, Priority Scheduling, SRTF, SJF, RR.', 'co': co2, 'btl': 'BTL3', 'marks': marks}
            ]
    
    if is_qp_type_2 and part_c_questions is None:
        co1 = 'CO1' if cia_type == 'CIA-I' else 'CO3'
        co2 = 'CO2' if cia_type == 'CIA-I' else 'CO4'
        part_c_questions = [
            {'qno': '9.(a)', 'question': 'Design and implement a priority scheduling algorithm. Analyze its performance with different process sets.', 'co': co1, 'btl': 'BTL4', 'marks': '16'},
            {'qno': '(OR)', 'question': '(OR)', 'co': '', 'btl': '', 'marks': ''},
            {'qno': '9.(b)', 'question': 'Evaluate the performance of different CPU scheduling algorithms for a given set of processes.', 'co': co2, 'btl': 'BTL5', 'marks': '16'}
        ]
    
    doc = Document()
    
    # ===== PAGE SETUP =====
    doc_section = doc.sections[0]
    doc_section.top_margin = Inches(0.4)
    doc_section.bottom_margin = Inches(0.5)
    doc_section.left_margin = Inches(0.6)
    doc_section.right_margin = Inches(0.6)
    
    # ===== FLOATING PAGE-ANCHORED LOGOS (No table, no text flow) =====
    # Logos are absolutely positioned, do not participate in document flow
    
    EMU_PER_INCH = 914400
    page_left_margin = 0.6 * EMU_PER_INCH
    page_right_edge = 8.5 * EMU_PER_INCH
    
    # Create anchor paragraph for floating logos
    logo_anchor_para = doc.add_paragraph()
    
    # Logo dimensions - SAME TOP ALIGNMENT (not equator)
    left_logo_width = 2.2  # inches - BIGGER
    left_logo_height = 2.1  # inches (approximate)
    right_logo_width = 1.0  # inches - SMALLER
    right_logo_height = 0.85  # inches (approximate)
    
    # EQUAL TOP POSITION for both logos (same vertical alignment)
    top_position = 0.25 * EMU_PER_INCH  # Same for both!
    
    # LEFT LOGO - M. Kumarasamy College (2.2 inches)
    college_logo_path = os.path.expanduser('~/Downloads/Kumarasamy.jpg')
    if os.path.exists(college_logo_path):
        left_logo_run = logo_anchor_para.add_run()
        left_logo_run.add_picture(college_logo_path, width=Inches(left_logo_width))
        try:
            # Make floating: page-anchored, In Front of Text
            make_image_floating(
                left_logo_run,
                width_emu=left_logo_width * EMU_PER_INCH,
                height_emu=left_logo_height * EMU_PER_INCH,
                pos_x_emu=page_left_margin,
                pos_v_emu=top_position  # SAME TOP
            )
        except Exception as e:
            print(f"Warning: Could not float left logo: {e}")
    
    # RIGHT LOGO - KR Group (1.0 inch)
    kr_logo_path = os.path.expanduser('~/Downloads/krl.jpg')
    if os.path.exists(kr_logo_path):
        right_logo_run = logo_anchor_para.add_run()
        right_logo_run.add_picture(kr_logo_path, width=Inches(right_logo_width))
        try:
            # Make floating: page-anchored, In Front of Text - SAME TOP POSITION
            make_image_floating(
                right_logo_run,
                width_emu=right_logo_width * EMU_PER_INCH,
                height_emu=right_logo_height * EMU_PER_INCH,
                pos_x_emu=page_right_edge - page_left_margin - (right_logo_width * EMU_PER_INCH),
                pos_v_emu=top_position  # SAME TOP AS LEFT
            )
        except Exception as e:
            print(f"Warning: Could not float right logo: {e}")

    
    doc.add_paragraph()  # Spacing
    
    # ===== CIA HEADING =====
    cia_label = 'CIA-1' if cia_type == 'CIA-I' else 'CIA-2'
    cia_para = doc.add_paragraph()
    cia_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    cia_run = cia_para.add_run(cia_label)
    cia_run.font.size = Pt(16)
    cia_run.font.bold = True
    
    doc.add_paragraph()
    
    # ===== REG NO TABLE (CENTERED) =====
    reg_table = doc.add_table(rows=1, cols=13)
    reg_table.style = 'Table Grid'
    reg_table.alignment = WD_TABLE_ALIGNMENT.CENTER  # CENTERED!
    reg_cells = reg_table.rows[0].cells
    
    reg_cells[0].text = 'REG\nNo'
    for para in reg_cells[0].paragraphs:
        para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        for run in para.runs:
            run.font.size = Pt(9)
            run.font.bold = True
    
    set_cell_width(reg_cells[0], 0.5)
    for i in range(1, 13):
        set_cell_width(reg_cells[i], 0.35)
    
    for cell in reg_cells:
        add_border_to_cell(cell)
    
    doc.add_paragraph()
    
    # ===== INFO TABLE (CENTERED) =====
    info_table = doc.add_table(rows=4, cols=4)
    info_table.style = 'Table Grid'
    info_table.alignment = WD_TABLE_ALIGNMENT.CENTER
    set_column_widths(info_table, [1.8, 1.9, 1.7, 1.6])
    
    # Row 1
    info_table.rows[0].cells[0].text = 'DEPARTMENT'
    info_table.rows[0].cells[1].text = department
    info_table.rows[0].cells[2].text = 'SEMESTER'
    info_table.rows[0].cells[3].text = semester
    
    # Row 2
    info_table.rows[1].cells[0].text = 'SECTION'
    info_table.rows[1].cells[1].text = section_val
    info_table.rows[1].cells[2].text = 'DATE & SESSION'
    info_table.rows[1].cells[3].text = date_session
    
    # Row 3
    info_table.rows[2].cells[0].text = 'DURATION'
    info_table.rows[2].cells[1].text = '120 Minutes'
    info_table.rows[2].cells[2].text = 'MAX MARKS'
    info_table.rows[2].cells[3].text = '60'
    
    # Row 4 - merged
    info_table.rows[3].cells[0].text = 'COURSE CODE & NAME'
    merged = info_table.rows[3].cells[1].merge(info_table.rows[3].cells[3])
    merged.text = f'{course_code}-{course_name}'
    
    # Style info table
    for row_idx, row in enumerate(info_table.rows):
        for col_idx, cell in enumerate(row.cells):
            add_border_to_cell(cell)
            for para in cell.paragraphs:
                para.alignment = WD_ALIGN_PARAGRAPH.CENTER
                for run in para.runs:
                    run.font.size = Pt(10)
                    if col_idx == 0 or col_idx == 2:
                        run.font.bold = True
    
    doc.add_paragraph()
    
    # ===== PART A =====
    part_a_heading = doc.add_paragraph()
    part_a_heading.alignment = WD_ALIGN_PARAGRAPH.CENTER
    part_a_run = part_a_heading.add_run('PART - A (6 X 2 MARKS = 12 MARKS)')
    part_a_run.font.size = Pt(11)
    part_a_run.font.bold = True
    
    headers = ['Q.NO', 'Questions', 'CO', 'BTL', 'MARKS']
    
    part_a_table = doc.add_table(rows=len(part_a_questions) + 1, cols=5)
    part_a_table.style = 'Table Grid'
    part_a_table.alignment = WD_TABLE_ALIGNMENT.CENTER
    set_column_widths(part_a_table, [0.5, 5.0, 0.45, 0.55, 0.6])
    
    for i, hdr in enumerate(headers):
        cell = part_a_table.rows[0].cells[i]
        cell.text = hdr
        for para in cell.paragraphs:
            para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for run in para.runs:
                run.font.bold = True
                run.font.size = Pt(10)
        add_border_to_cell(cell)
    
    for i, q in enumerate(part_a_questions):
        row_data = [q.get('qno', ''), q.get('question', ''), q.get('co', ''), q.get('btl', ''), q.get('marks', '')]
        for j, text in enumerate(row_data):
            cell = part_a_table.rows[i+1].cells[j]
            cell.text = str(text)
            for para in cell.paragraphs:
                para.alignment = WD_ALIGN_PARAGRAPH.CENTER if j != 1 else WD_ALIGN_PARAGRAPH.LEFT
                for run in para.runs:
                    run.font.size = Pt(10)
            add_border_to_cell(cell)
    
    doc.add_paragraph()
    
    # ===== PART B =====
    part_b_heading_text = f'PART - B ({part_b_count} X {part_b_marks_each} MARKS = {part_b_count * part_b_marks_each} MARKS)'
    part_b_heading = doc.add_paragraph()
    part_b_heading.alignment = WD_ALIGN_PARAGRAPH.CENTER
    part_b_run = part_b_heading.add_run(part_b_heading_text)
    part_b_run.font.size = Pt(11)
    part_b_run.font.bold = True
    
    part_b_table = doc.add_table(rows=len(part_b_questions) + 1, cols=5)
    part_b_table.style = 'Table Grid'
    part_b_table.alignment = WD_TABLE_ALIGNMENT.CENTER
    set_column_widths(part_b_table, [0.5, 5.0, 0.45, 0.55, 0.6])
    
    for i, hdr in enumerate(headers):
        cell = part_b_table.rows[0].cells[i]
        cell.text = hdr
        for para in cell.paragraphs:
            para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for run in para.runs:
                run.font.bold = True
                run.font.size = Pt(10)
        add_border_to_cell(cell)
    
    for i, q in enumerate(part_b_questions):
        row_data = [q.get('qno', ''), q.get('question', ''), q.get('co', ''), q.get('btl', ''), q.get('marks', '')]
        is_or_row = row_data[0] == '(OR)' or row_data[1] == '(OR)'
        for j, text in enumerate(row_data):
            cell = part_b_table.rows[i+1].cells[j]
            cell.text = str(text)
            for para in cell.paragraphs:
                para.alignment = WD_ALIGN_PARAGRAPH.CENTER if (is_or_row or j != 1) else WD_ALIGN_PARAGRAPH.LEFT
                for run in para.runs:
                    run.font.size = Pt(10)
                    if is_or_row:
                        run.font.bold = True
            add_border_to_cell(cell)
    
    # ===== PART C (QP-II only) =====
    if is_qp_type_2 and part_c_questions:
        doc.add_paragraph()
        
        part_c_heading = doc.add_paragraph()
        part_c_heading.alignment = WD_ALIGN_PARAGRAPH.CENTER
        part_c_run = part_c_heading.add_run('PART - C (1 X 16 MARKS = 16 MARKS)')
        part_c_run.font.size = Pt(11)
        part_c_run.font.bold = True
        
        part_c_table = doc.add_table(rows=len(part_c_questions) + 1, cols=5)
        part_c_table.style = 'Table Grid'
        part_c_table.alignment = WD_TABLE_ALIGNMENT.CENTER
        set_column_widths(part_c_table, [0.5, 5.0, 0.45, 0.55, 0.6])
        
        for i, hdr in enumerate(headers):
            cell = part_c_table.rows[0].cells[i]
            cell.text = hdr
            for para in cell.paragraphs:
                para.alignment = WD_ALIGN_PARAGRAPH.CENTER
                for run in para.runs:
                    run.font.bold = True
                    run.font.size = Pt(10)
            add_border_to_cell(cell)
        
        for i, q in enumerate(part_c_questions):
            row_data = [q.get('qno', ''), q.get('question', ''), q.get('co', ''), q.get('btl', ''), q.get('marks', '')]
            is_or_row = row_data[0] == '(OR)' or row_data[1] == '(OR)'
            for j, text in enumerate(row_data):
                cell = part_c_table.rows[i+1].cells[j]
                cell.text = str(text)
                for para in cell.paragraphs:
                    para.alignment = WD_ALIGN_PARAGRAPH.CENTER if (is_or_row or j != 1) else WD_ALIGN_PARAGRAPH.LEFT
                    for run in para.runs:
                        run.font.size = Pt(10)
                        if is_or_row:
                            run.font.bold = True
                add_border_to_cell(cell)
    
    # Save or return bytes
    if output_path:
        doc.save(output_path)
        return output_path
    else:
        doc_bytes = io.BytesIO()
        doc.save(doc_bytes)
        doc_bytes.seek(0)
        return doc_bytes.getvalue()


if __name__ == "__main__":
    try:
        create_exam_paper(output_path='CIA_Exam_Paper_QP1.docx', cia_type='CIA-I', qp_type='QP-I')
        print("✅ CIA-I QP-I created: CIA_Exam_Paper_QP1.docx")
        
        create_exam_paper(output_path='CIA_Exam_Paper_QP2.docx', cia_type='CIA-II', qp_type='QP-II')
        print("✅ CIA-II QP-II created: CIA_Exam_Paper_QP2.docx")
        
        print(f"📁 Saved in: {os.getcwd()}")
    except Exception as e:
        print(f"❌ Error: {e}")
        import traceback
        traceback.print_exc()
